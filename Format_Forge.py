### main.py
import sys
from core.document_processor import process_document

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python main.py archivo.docx [formato]")
        sys.exit(1)

    filepath = sys.argv[1]
    output_format = sys.argv[2] if len(sys.argv) > 2 else "docx"
    process_document(filepath, output_format)


### core/document_processor.py
from ..utils.docx_loader import load_docx
from ..styles.apa_formatter import apply_apa_style
from ..output.exporter import export_doc
from ..utils.section_detector import detect_sections

def process_document(filepath, output_format="docx"):
    document = load_docx(filepath)
    sections = detect_sections(document)
    formatted_document = apply_apa_style(document)
    export_doc(formatted_document, filepath, output_format)


### utils/docx_loader.py
from docx import Document

def load_docx(filepath):
    return Document(filepath)


### styles/apa_formatter.py
from ..utils.section_detector import detect_sections
from docx.shared import Pt
import re

def is_title_level(text):
    if re.match(r"^(\d+\.)+\s+.*", text):
        level = text.count('.')
        return min(level, 3)
    return None

def apply_apa_style(document):
    sections = detect_sections(document)
    for i, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        paragraph.style = 'Normal'
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.paragraph_format.line_spacing = 2.0
        paragraph.paragraph_format.left_indent = Pt(0)

        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        if i in sections.get('title', []):
            paragraph.alignment = 1
            for run in paragraph.runs:
                run.bold = True
        elif i in sections.get('abstract', []):
            paragraph.paragraph_format.left_indent = Pt(0.5 * 72)
        elif i in sections.get('references', []):
            paragraph.paragraph_format.first_line_indent = Pt(-0.5 * 72)
            paragraph.paragraph_format.left_indent = Pt(0.5 * 72)

        title_level = is_title_level(text)
        if title_level == 1:
            paragraph.alignment = 1
            for run in paragraph.runs:
                run.bold = True
                run.italic = False
        elif title_level == 2:
            paragraph.alignment = 0
            for run in paragraph.runs:
                run.bold = True
                run.italic = False
        elif title_level == 3:
            paragraph.alignment = 0
            for run in paragraph.runs:
                run.bold = False
                run.italic = True

    return document


### utils/section_detector.py
import re

SECTION_KEYWORDS = {
    "title": ["título", "title"],
    "abstract": ["resumen", "abstract"],
    "introduction": ["introducción", "introduction"],
    "conclusion": ["conclusión", "conclusion", "concluding"],
    "references": ["referencias", "bibliografía", "references"],
}

def detect_sections(document):
    detected = {key: [] for key in SECTION_KEYWORDS}
    for i, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip().lower()
        for section, keywords in SECTION_KEYWORDS.items():
            if any(re.fullmatch(rf"{kw}[:\s]*", text) for kw in keywords):
                detected[section].append(i)
    return detected


### output/exporter.py
import os
from .pdf_exporter import export_to_pdf

def export_doc(document, original_path, output_format):
    base, _ = os.path.splitext(original_path)
    docx_path = base + "_APA.docx"
    document.save(docx_path)

    if output_format == "pdf":
        return export_to_pdf(docx_path)
    return docx_path


### output/pdf_exporter.py
import os
import platform
import subprocess

def export_to_pdf(docx_path):
    system = platform.system()
    if system == "Windows":
        try:
            from docx2pdf import convert
            convert(docx_path)
            return docx_path.replace(".docx", ".pdf")
        except ImportError:
            raise ImportError("Necesitas instalar docx2pdf: pip install docx2pdf")
    else:
        output_path = docx_path.replace(".docx", ".pdf")
        command = ["libreoffice", "--headless", "--convert-to", "pdf", docx_path, "--outdir", os.path.dirname(docx_path)]
        result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if result.returncode == 0:
            return output_path
        else:
            raise RuntimeError("Falló la conversión a PDF usando LibreOffice")
